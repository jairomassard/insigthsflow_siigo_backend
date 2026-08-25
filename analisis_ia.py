"""
analisis_ia.py
================
Motor del feature "Analizar con IA" para reportes financieros. Arranca
con Estado de Resultados (PyG); pensado para reusarse con otros reportes
mas adelante siguiendo el mismo patron de cache/tope/huella.

Decisiones de diseno (ver memoria
project_insightflow_analisis_ia_reportes para el detalle completo y el
por que de cada una):

- Modelo: Opus 5. Elegido tras comparar calidad real (no teorica) contra
  Sonnet 5, corriendo el mismo prompt contra el PyG real de un cliente
  (Binaria Media). Opus hizo cruces de datos que Sonnet no hizo.
- Tope: 15 analisis REALES por mes por cliente. Los hits de cache NO
  cuentan contra el tope.
- Cache por huella de datos, no solo por periodo: los numeros de un
  periodo ya cerrado en InsightFlow SI pueden cambiar despues (resync,
  correccion de retencion, anulacion de factura, backfill). Si la
  huella de los datos actuales no coincide con la guardada, se
  regenera - nunca se sirve un analisis basado en cifras obsoletas.
- Produccion NO le pide al cliente exportar nada - se le manda a la IA
  el JSON que el backend ya calcula para el reporte en pantalla
  (construir_pnl), curado a los campos relevantes, no un dump crudo.
- Prompt acotado a un formato de largo fijo (~1200 palabras) - los
  analisis "a fondo" sin limite salieron de 10-16K tokens de salida en
  las pruebas, demasiado para un modal y mas caro de lo necesario.
"""

import os
import re
import json
import hashlib
from datetime import date, datetime
from io import BytesIO

import anthropic
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")  # sin display - backend headless, solo genera PNG
import matplotlib.pyplot as plt

# Mismos colores que el grafico "Tendencia P&L Mensual" en pantalla
# (recharts, frontend/.../estado-resultados/page.tsx) - para que el
# grafico del Word se vea consistente con lo que el cliente ya vio.
COLOR_INGRESOS = "#10b981"
COLOR_COSTOS_GASTOS = "#f43f5e"
COLOR_EBITDA = "#4f46e5"

from models import db, AnalisisIAUso, AnalisisIACache


MODELO_ANALISIS = "claude-opus-5"
TOPE_MENSUAL = 15

# Referencial, solo para guardar un estimado de costo en COP junto al
# analisis (metricas internas) - no se usa para facturarle nada al
# cliente, el feature se regala dentro del plan Completo.
TRM_COP_POR_USD = 3063

PRECIO_POR_MILLON_USD = {
    "claude-opus-5": {"input": 5.00, "output": 25.00},
}

PROMPT_SISTEMA = (
    "Eres un analista financiero senior ayudando al dueño o gerente de una "
    "PyME colombiana a entender su Estado de Resultados (PyG). Vas a "
    "recibir datos ya calculados: KPIs del periodo, evolución mes a mes, "
    "y composición por cuenta contable - no inventes cifras que no estén "
    "ahí. Responde siempre en español, en formato markdown, con esta "
    "estructura fija y sin exceder ~2500-3000 palabras en total (podés "
    "usar tablas markdown donde ayuden a mostrar un desglose, ej. la "
    "cascada del PyG o una comparación mes a mes - las tablas no cuentan "
    "tan estricto contra el límite de palabras como el texto corrido):\n\n"
    "## Resumen ejecutivo\n"
    "2-3 frases con el diagnóstico central - qué es lo que más está "
    "afectando la utilidad.\n\n"
    "## Hallazgos principales\n"
    "Entre 3 y 5 hallazgos, cada uno con el número que lo respalda y por "
    "qué importa. Prioriza por impacto en la utilidad, no por orden "
    "cronológico ni por orden de aparición en los datos. Usa tablas donde "
    "el desglose lo justifique.\n\n"
    "## Recomendaciones\n"
    "Entre 3 y 5 acciones concretas y accionables, ligadas directamente a "
    "los hallazgos de arriba.\n\n"
    "Si algo no se puede determinar con los datos disponibles (por "
    "ejemplo, si no hay una línea de impuesto de renta), dilo "
    "explícitamente en vez de asumir o de omitirlo."
)


PROMPT_SISTEMA_BALANCE = (
    "Eres un analista financiero senior ayudando al dueño o gerente de una "
    "PyME colombiana a entender su Estado de Situación Financiera (Balance "
    "General). Vas a recibir datos ya calculados en JSON: los indicadores "
    "clave del corte (liquidez, endeudamiento, autonomía financiera, "
    "capital de trabajo, cuadratura), el detalle de patrimonio explícito "
    "vs. calculado, alertas de calidad de datos ya detectadas por el "
    "sistema, y el balance agrupado por cuenta contable (activo corriente, "
    "activo no corriente, pasivo corriente, pasivo no corriente y "
    "patrimonio) con su saldo actual y, si el usuario pidió comparación, "
    "el saldo del corte anterior y su variación. No inventes cifras que no "
    "estén ahí. Responde siempre en español, en formato markdown, con esta "
    "estructura fija y sin exceder ~2500-3000 palabras en total (podés "
    "usar tablas markdown donde ayuden a mostrar un desglose):\n\n"
    "## Resumen ejecutivo\n"
    "2-3 frases con el diagnóstico central: cómo está la empresa en "
    "liquidez, endeudamiento y patrimonio, y cuál es el hallazgo más "
    "importante.\n\n"
    "## Hallazgos principales\n"
    "Entre 3 y 6 hallazgos, cada uno con el número que lo respalda y por "
    "qué importa. IMPORTANTE sobre la comparación: cuando los ítems "
    "traigan saldo_anterior y variacion_abs (el usuario pidió comparar "
    "contra otro corte), no la trates como un tema aparte al final - "
    "usala DENTRO de cada hallazgo relevante para mostrar trayectoria, no "
    "solo posición (ej. 'la razón corriente es 1.4, y bajó desde 1.8 en "
    "el corte anterior', no solo 'la razón corriente es 1.4'). Cuando NO "
    "haya comparación, cubrí los mismos puntos solo con el corte actual - "
    "el análisis debe ser igual de completo en ambos casos. Cubrí "
    "explícitamente, cuando los datos lo permitan:\n"
    "- Liquidez: si la razón corriente (kpis.razon_corriente) y el "
    "capital de trabajo (kpis.capital_trabajo) alcanzan para cubrir las "
    "obligaciones de corto plazo, y si hay comparación, si esa cobertura "
    "mejoró o empeoró frente al corte anterior.\n"
    "- Endeudamiento y autonomía financiera (kpis.nivel_endeudamiento_pct, "
    "kpis.autonomia_financiera_pct): qué tan dependiente es la empresa de "
    "terceros vs. patrimonio propio, si es un nivel sano o de riesgo, y si "
    "hay comparación, hacia dónde se está moviendo esa dependencia.\n"
    "- Composición y concentración: si hay una sola cuenta que concentra "
    "una parte desproporcionada de un grupo (por ejemplo, una cuenta "
    "'Otros' con la mitad de un pasivo) - eso casi siempre amerita "
    "revisión, no es solo un número. Si hay comparación, señalá también "
    "las cuentas o grupos con el mayor crecimiento o caída en $ o % entre "
    "ambos cortes, no solo la composición del corte actual.\n"
    "- Calidad del patrimonio: revisá patrimonio_detalle - si "
    "usa_patrimonio_calculado es true, o si patrimonio_calculado_total "
    "representa una porción grande de patrimonio_total frente a "
    "patrimonio_explicito_total, decilo con claridad - es señal de que la "
    "contabilidad no está cerrando el ejercicio correctamente (falta "
    "registrar o trasladar la utilidad del período), no un tecnicismo "
    "menor.\n"
    "- Cuadratura y alertas de calidad de datos: si kpis.cuadratura es "
    "distinta de cero, o si viene algo en alertas_calidad_datos, "
    "repórtalo como hallazgo prioritario (primero en la lista) - "
    "compromete la confiabilidad de todo lo demás.\n\n"
    "## Recomendaciones\n"
    "Entre 3 y 5 acciones concretas y accionables, ligadas directamente a "
    "los hallazgos de arriba.\n\n"
    "Si algo no se puede determinar con los datos disponibles, dilo "
    "explícitamente en vez de asumir u omitirlo."
)


PROMPT_SISTEMA_INDICADORES = (
    "Eres un analista financiero senior ayudando al dueño o gerente de una "
    "PyME colombiana a entender su reporte de Indicadores Financieros. "
    "Este reporte combina el Balance General al corte final (liquidez, "
    "endeudamiento, patrimonio) con el Estado de Resultados del período "
    "seleccionado (ingresos, costos, utilidad), más indicadores cruzados "
    "que combinan ambos (ROE, ROA, prueba ácida, días de cobro/pago "
    "reales usando cartera y cuentas por pagar reales, cobertura de "
    "intereses). Vas a recibir todo eso ya calculado en JSON - no "
    "inventes cifras que no estén ahí.\n\n"
    "MUY IMPORTANTE sobre los parámetros de la empresa: revisá el campo "
    "parametros_configurados.\n"
    "- Si es false (modo informativo): la empresa NO ha definido metas "
    "propias todavía. Aun así tenés que entregar un análisis COMPLETO y "
    "útil - explicá cómo está la empresa hoy (liquidez, endeudamiento, "
    "rentabilidad, ciclo de caja) usando criterio financiero general, sin "
    "inventar ni asumir metas que la empresa no definió, y sin frases "
    "tipo 'está mal' o 'está fuera de rango' (no hay rango propio contra "
    "qué medir). Cerrá sugiriendo, si tiene sentido, que configure sus "
    "propios parámetros para activar la comparación contra metas.\n"
    "- Si es true (modo parametrizado): además de lo anterior, usá "
    "interpretaciones_detalle y config_financiera para decir "
    "explícitamente qué indicadores están dentro o fuera de la meta que "
    "la empresa misma definió, y qué tan lejos. Pero esto es una CAPA "
    "ADICIONAL, no un reemplazo del análisis conjunto - seguí explicando "
    "el panorama completo del negocio, no te limites a listar qué "
    "indicadores pasaron o no la meta.\n\n"
    "En ambos modos, el análisis debe ser una lectura CONJUNTA de todo el "
    "reporte, no indicador por indicador aislado. Cruzá los datos: por "
    "ejemplo, ROE alto pero ROA bajo indica que la rentabilidad viene del "
    "apalancamiento y no de la eficiencia del activo; DSO mucho mayor que "
    "DPO indica que la empresa está financiando a sus clientes con el "
    "dinero que le debe a sus proveedores. Usá también resumen_balance "
    "(narrativa y alertas ya calculadas del balance) y meta_balance "
    "(patrimonio explícito vs. calculado, activo no corriente bruto vs. "
    "neto) como parte del panorama - no los ignores, son parte de lo que "
    "el usuario ve en este mismo reporte.\n\n"
    "Responde siempre en español, en formato markdown, con esta "
    "estructura fija y sin exceder ~2500-3000 palabras en total (podés "
    "usar tablas markdown donde ayuden a mostrar un desglose):\n\n"
    "## Resumen ejecutivo\n"
    "2-3 frases: cómo está la empresa en conjunto (liquidez, "
    "endeudamiento, rentabilidad) y cuál es el hallazgo más importante.\n"
    "\n"
    "## Hallazgos principales\n"
    "Entre 4 y 7 hallazgos, cada uno con el número que lo respalda y por "
    "qué importa. Cubrí, cuando los datos lo permitan:\n"
    "- Liquidez y ciclo de caja: razón corriente (kpis liquidez), capital "
    "de trabajo, prueba ácida, y la relación entre DSO y DPO (¿quién "
    "está financiando a quién?).\n"
    "- Endeudamiento y estructura: apalancamiento, autonomía financiera, "
    "endeudamiento de largo plazo.\n"
    "- Rentabilidad: margen neto (rentabilidad), ROE, ROA, y qué explica "
    "la diferencia entre ellos si es significativa.\n"
    "- Cobertura de deuda: si hay gastos financieros, si la utilidad "
    "operativa alcanza para cubrirlos cómodamente (cobertura_intereses).\n"
    "- Calidad del balance detrás de los indicadores: si meta_balance "
    "muestra que buena parte del patrimonio es calculado (no explícito), "
    "o hay alertas en resumen_balance, decilo con claridad - afecta qué "
    "tan confiables son los indicadores que dependen del patrimonio "
    "(autonomía, endeudamiento largo plazo, ROE).\n"
    "- Si parametros_configurados es true, qué indicadores están fuera "
    "de la meta propia de la empresa y qué tan lejos.\n\n"
    "## Recomendaciones\n"
    "Entre 3 y 5 acciones concretas y accionables, ligadas directamente a "
    "los hallazgos de arriba.\n\n"
    "Si algo no se puede determinar con los datos disponibles, dilo "
    "explícitamente en vez de asumir u omitirlo."
)


class TopeAlcanzadoError(Exception):
    """El cliente ya usó sus análisis reales del mes."""

    def __init__(self, uso_actual: int, tope: int):
        self.uso_actual = uso_actual
        self.tope = tope
        super().__init__(f"Tope mensual de análisis con IA alcanzado ({uso_actual}/{tope}).")


def _huella_datos(payload: dict) -> str:
    canon = json.dumps(payload, sort_keys=True, default=str, ensure_ascii=False)
    return hashlib.sha256(canon.encode("utf-8")).hexdigest()


def _periodo_actual() -> tuple[int, int]:
    hoy = date.today()
    return hoy.year, hoy.month


def consultar_uso_mensual(idcliente: int) -> int:
    anio, mes = _periodo_actual()
    row = AnalisisIAUso.query.filter_by(
        idcliente=idcliente, periodo_anio=anio, periodo_mes=mes
    ).first()
    return row.cantidad if row else 0


def _incrementar_uso_mensual(idcliente: int) -> None:
    anio, mes = _periodo_actual()
    row = AnalisisIAUso.query.filter_by(
        idcliente=idcliente, periodo_anio=anio, periodo_mes=mes
    ).first()
    if row:
        row.cantidad += 1
    else:
        row = AnalisisIAUso(idcliente=idcliente, periodo_anio=anio, periodo_mes=mes, cantidad=1)
        db.session.add(row)
    db.session.commit()


def _resumen_pyg_para_ia(pnl_data: dict) -> dict:
    """Payload curado que se le manda a la IA - los 3 bloques que ya
    calcula construir_pnl(), no un dump crudo de toda la contabilidad."""
    return {
        "kpis_totales_periodo": pnl_data.get("kpis", {}),
        "evolucion_mensual": pnl_data.get("evolucion", []),
        "composicion_por_cuenta": pnl_data.get("composicion", []),
    }


def _resultado_verificacion(cache, huella: str) -> dict:
    """Respuesta liviana para el chequeo "¿esto va a salir gratis del
    caché o va a generar un análisis nuevo?" - SIN llamar a la IA. Sirve
    para avisarle al usuario ANTES de gastar cupo, no solo después."""
    return {
        "actualizado": bool(cache and cache.huella_datos == huella),
        "existe_cache": cache is not None,
    }


def generar_analisis_pyg(
    idcliente: int,
    desde: str,
    hasta: str,
    pnl_data: dict,
    forzar: bool = False,
    solo_verificar: bool = False,
) -> dict:
    """
    Punto de entrada principal para el Estado de Resultados.

    Devuelve un dict con "fuente" ("cache" o "nuevo"), "analisis"
    (markdown), "modelo" y "generado_en". Levanta TopeAlcanzadoError si
    hace falta generar de nuevo pero el cliente ya no tiene cupo este mes.

    solo_verificar=True corta antes de llamar a la IA y devuelve solo si
    el caché sigue vigente - usado por el endpoint de "verificar" para
    poder avisarle al usuario antes de gastar, no después.
    """
    tipo_reporte = "estado_resultados"
    periodo_desde = date.fromisoformat(desde)
    periodo_hasta = date.fromisoformat(hasta)

    payload = _resumen_pyg_para_ia(pnl_data)
    huella = _huella_datos(payload)

    cache = AnalisisIACache.query.filter_by(
        idcliente=idcliente,
        tipo_reporte=tipo_reporte,
        periodo_desde=periodo_desde,
        periodo_hasta=periodo_hasta,
    ).first()

    if solo_verificar:
        return _resultado_verificacion(cache, huella)

    if cache and cache.huella_datos == huella and not forzar:
        return {
            "fuente": "cache",
            "analisis": cache.analisis_texto,
            "modelo": cache.modelo,
            "generado_en": cache.updated_at.isoformat() if cache.updated_at else None,
            "uso_mensual": consultar_uso_mensual(idcliente),
            "tope_mensual": TOPE_MENSUAL,
        }

    uso_actual = consultar_uso_mensual(idcliente)
    if uso_actual >= TOPE_MENSUAL:
        raise TopeAlcanzadoError(uso_actual, TOPE_MENSUAL)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY no está configurada en el entorno del backend.")

    client = anthropic.Anthropic(api_key=api_key)

    mensaje_usuario = (
        f"Aquí está el Estado de Resultados del período {desde} a {hasta}, "
        "en formato JSON:\n\n"
        f"{json.dumps(payload, ensure_ascii=False, default=str)}"
    )

    # max_tokens generoso a propósito y con streaming: con thinking
    # adaptativo el razonamiento consume del mismo cupo de salida antes de
    # escribir la respuesta final - un tope justo corta la respuesta a la
    # mitad (visto DOS VECES en pruebas reales de esta sesión, primero con
    # 4000 y de nuevo con 10000). Streaming evita además el guard de la
    # SDK que rechaza max_tokens grandes sin streaming.
    with client.beta.messages.stream(
        model=MODELO_ANALISIS,
        max_tokens=32000,
        system=PROMPT_SISTEMA,
        thinking={"type": "adaptive"},
        betas=["server-side-fallback-2026-07-01"],
        fallbacks="default",
        messages=[{"role": "user", "content": mensaje_usuario}],
    ) as stream:
        response = stream.get_final_message()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            "La respuesta de la IA se cortó por max_tokens antes de terminar. "
            "No se guardó en caché ni se consumió el cupo mensual - intenta de nuevo."
        )

    texto = "".join(block.text for block in response.content if block.type == "text")

    tokens_in = response.usage.input_tokens
    tokens_out = response.usage.output_tokens
    precio = PRECIO_POR_MILLON_USD[MODELO_ANALISIS]
    costo_usd = (tokens_in / 1_000_000 * precio["input"]) + (tokens_out / 1_000_000 * precio["output"])
    costo_cop = round(costo_usd * TRM_COP_POR_USD, 2)

    if cache:
        cache.huella_datos = huella
        cache.modelo = MODELO_ANALISIS
        cache.analisis_texto = texto
        cache.tokens_entrada = tokens_in
        cache.tokens_salida = tokens_out
        cache.costo_cop = costo_cop
    else:
        cache = AnalisisIACache(
            idcliente=idcliente,
            tipo_reporte=tipo_reporte,
            periodo_desde=periodo_desde,
            periodo_hasta=periodo_hasta,
            huella_datos=huella,
            modelo=MODELO_ANALISIS,
            analisis_texto=texto,
            tokens_entrada=tokens_in,
            tokens_salida=tokens_out,
            costo_cop=costo_cop,
        )
        db.session.add(cache)

    db.session.commit()
    _incrementar_uso_mensual(idcliente)

    return {
        "fuente": "nuevo",
        "analisis": texto,
        "modelo": MODELO_ANALISIS,
        "generado_en": datetime.utcnow().isoformat(),
        "uso_mensual": consultar_uso_mensual(idcliente),
        "tope_mensual": TOPE_MENSUAL,
    }


def _resumen_balance_para_ia(balance_data: dict) -> dict:
    """Payload curado para el Balance General - usa exactamente las mismas
    listas (balance.activo_corriente, balance.patrimonio, etc.) que
    consume la cascada visual en pantalla, para que la IA lea lo mismo
    que ve el usuario. Incluye ademas el detalle de patrimonio explicito
    vs calculado y las alertas de calidad de datos que el sistema ya
    calcula (resumen.alertas) - son la senal mas confiable para detectar
    problemas de cierre contable, no hay que pedirle a la IA que las
    re-derive desde cero."""
    balance = balance_data.get("balance", {})
    return {
        "fechas": balance_data.get("fechas", {}),
        "kpis": balance_data.get("kpis", {}),
        "patrimonio_detalle": balance_data.get("meta", {}).get("patrimonio", {}),
        "alertas_calidad_datos": balance_data.get("resumen", {}).get("alertas", []),
        "activo_corriente": balance.get("activo_corriente", []),
        "activo_no_corriente": balance.get("activo_no_corriente", []),
        "pasivo_corriente": balance.get("pasivo_corriente", []),
        "pasivo_no_corriente": balance.get("pasivo_no_corriente", []),
        "patrimonio": balance.get("patrimonio", []),
    }


def generar_analisis_balance(
    idcliente: int,
    fecha_corte: str,
    comparar_con: str | None,
    balance_data: dict,
    forzar: bool = False,
    solo_verificar: bool = False,
) -> dict:
    """
    Punto de entrada principal para el Balance General. Mismo patron de
    cache/tope/huella que generar_analisis_pyg().

    El Balance no maneja un rango desde/hasta como el PyG, sino un corte
    puntual (y opcionalmente otro corte de comparacion) - se reusan las
    columnas periodo_desde/periodo_hasta de AnalisisIACache sin migracion:
    periodo_hasta guarda el corte que se esta analizando y periodo_desde
    guarda el corte de comparacion (o el mismo corte si no hay
    comparacion), asi cada combinacion de cortes cachea por separado.

    solo_verificar=True corta antes de llamar a la IA (ver
    generar_analisis_pyg).
    """
    tipo_reporte = "balance_general"
    periodo_hasta = date.fromisoformat(fecha_corte)
    periodo_desde = date.fromisoformat(comparar_con) if comparar_con else periodo_hasta

    payload = _resumen_balance_para_ia(balance_data)
    huella = _huella_datos(payload)

    cache = AnalisisIACache.query.filter_by(
        idcliente=idcliente,
        tipo_reporte=tipo_reporte,
        periodo_desde=periodo_desde,
        periodo_hasta=periodo_hasta,
    ).first()

    if solo_verificar:
        return _resultado_verificacion(cache, huella)

    if cache and cache.huella_datos == huella and not forzar:
        return {
            "fuente": "cache",
            "analisis": cache.analisis_texto,
            "modelo": cache.modelo,
            "generado_en": cache.updated_at.isoformat() if cache.updated_at else None,
            "uso_mensual": consultar_uso_mensual(idcliente),
            "tope_mensual": TOPE_MENSUAL,
        }

    uso_actual = consultar_uso_mensual(idcliente)
    if uso_actual >= TOPE_MENSUAL:
        raise TopeAlcanzadoError(uso_actual, TOPE_MENSUAL)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY no está configurada en el entorno del backend.")

    client = anthropic.Anthropic(api_key=api_key)

    mensaje_usuario = (
        f"Aquí está el Balance General con corte al {fecha_corte}"
        + (f", comparado con el corte al {comparar_con}" if comparar_con else "")
        + ", en formato JSON:\n\n"
        + json.dumps(payload, ensure_ascii=False, default=str)
    )

    with client.beta.messages.stream(
        model=MODELO_ANALISIS,
        max_tokens=32000,
        system=PROMPT_SISTEMA_BALANCE,
        thinking={"type": "adaptive"},
        betas=["server-side-fallback-2026-07-01"],
        fallbacks="default",
        messages=[{"role": "user", "content": mensaje_usuario}],
    ) as stream:
        response = stream.get_final_message()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            "La respuesta de la IA se cortó por max_tokens antes de terminar. "
            "No se guardó en caché ni se consumió el cupo mensual - intenta de nuevo."
        )

    texto = "".join(block.text for block in response.content if block.type == "text")

    tokens_in = response.usage.input_tokens
    tokens_out = response.usage.output_tokens
    precio = PRECIO_POR_MILLON_USD[MODELO_ANALISIS]
    costo_usd = (tokens_in / 1_000_000 * precio["input"]) + (tokens_out / 1_000_000 * precio["output"])
    costo_cop = round(costo_usd * TRM_COP_POR_USD, 2)

    if cache:
        cache.huella_datos = huella
        cache.modelo = MODELO_ANALISIS
        cache.analisis_texto = texto
        cache.tokens_entrada = tokens_in
        cache.tokens_salida = tokens_out
        cache.costo_cop = costo_cop
    else:
        cache = AnalisisIACache(
            idcliente=idcliente,
            tipo_reporte=tipo_reporte,
            periodo_desde=periodo_desde,
            periodo_hasta=periodo_hasta,
            huella_datos=huella,
            modelo=MODELO_ANALISIS,
            analisis_texto=texto,
            tokens_entrada=tokens_in,
            tokens_salida=tokens_out,
            costo_cop=costo_cop,
        )
        db.session.add(cache)

    db.session.commit()
    _incrementar_uso_mensual(idcliente)

    return {
        "fuente": "nuevo",
        "analisis": texto,
        "modelo": MODELO_ANALISIS,
        "generado_en": datetime.utcnow().isoformat(),
        "uso_mensual": consultar_uso_mensual(idcliente),
        "tope_mensual": TOPE_MENSUAL,
    }


def _resumen_indicadores_para_ia(data: dict) -> dict:
    """Payload curado para Indicadores Financieros - la respuesta cruda
    del endpoint trae ademas resumen_financiero (una tabla redundante con
    "indicadores", solo renombrada) que no aporta nada nuevo a la IA, por
    eso no se incluye aca. Se manda explicaciones solo de las formulas
    "no obvias" (las que cruzan balance con P&L, ej. DSO/DPO usan cartera
    y CxP reales, no el balance) para evitar que la IA las confunda con
    una version simplificada del ratio."""
    formulas_no_obvias = ("dso_dias_cobro", "dpo_dias_pago", "prueba_acida", "cobertura_intereses", "roe", "roa")
    explicaciones = data.get("explicaciones") or {}

    return {
        "meta": data.get("meta", {}),
        "parametros_configurados": bool(data.get("parametros_configurados")),
        "config_financiera": data.get("config_financiera") if data.get("parametros_configurados") else None,
        "indicadores": data.get("indicadores", {}),
        "formulas_indicadores_cruzados": {k: explicaciones[k] for k in formulas_no_obvias if k in explicaciones},
        "interpretaciones_detalle": data.get("interpretaciones_detalle", {}),
        "conclusiones_parametrizadas": data.get("conclusiones", []),
        "evolucion_mensual": data.get("evolucion_mensual", []),
        "resumen_balance": data.get("resumen_balance", {}),
        "meta_balance": data.get("meta_balance", {}),
    }


def generar_analisis_indicadores(
    idcliente: int,
    fecha_desde: str,
    fecha_hasta: str,
    indicadores_data: dict,
    forzar: bool = False,
    solo_verificar: bool = False,
) -> dict:
    """
    Punto de entrada principal para Indicadores Financieros. Mismo patron
    de cache/tope/huella que generar_analisis_pyg()/generar_analisis_balance().

    solo_verificar=True corta antes de llamar a la IA (ver
    generar_analisis_pyg).
    """
    tipo_reporte = "indicadores_financieros"
    periodo_desde = date.fromisoformat(fecha_desde)
    periodo_hasta = date.fromisoformat(fecha_hasta)

    payload = _resumen_indicadores_para_ia(indicadores_data)
    huella = _huella_datos(payload)

    cache = AnalisisIACache.query.filter_by(
        idcliente=idcliente,
        tipo_reporte=tipo_reporte,
        periodo_desde=periodo_desde,
        periodo_hasta=periodo_hasta,
    ).first()

    if solo_verificar:
        return _resultado_verificacion(cache, huella)

    if cache and cache.huella_datos == huella and not forzar:
        return {
            "fuente": "cache",
            "analisis": cache.analisis_texto,
            "modelo": cache.modelo,
            "generado_en": cache.updated_at.isoformat() if cache.updated_at else None,
            "uso_mensual": consultar_uso_mensual(idcliente),
            "tope_mensual": TOPE_MENSUAL,
        }

    uso_actual = consultar_uso_mensual(idcliente)
    if uso_actual >= TOPE_MENSUAL:
        raise TopeAlcanzadoError(uso_actual, TOPE_MENSUAL)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY no está configurada en el entorno del backend.")

    client = anthropic.Anthropic(api_key=api_key)

    mensaje_usuario = (
        f"Aquí está el reporte de Indicadores Financieros del período "
        f"{fecha_desde} a {fecha_hasta}, en formato JSON:\n\n"
        f"{json.dumps(payload, ensure_ascii=False, default=str)}"
    )

    with client.beta.messages.stream(
        model=MODELO_ANALISIS,
        max_tokens=32000,
        system=PROMPT_SISTEMA_INDICADORES,
        thinking={"type": "adaptive"},
        betas=["server-side-fallback-2026-07-01"],
        fallbacks="default",
        messages=[{"role": "user", "content": mensaje_usuario}],
    ) as stream:
        response = stream.get_final_message()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            "La respuesta de la IA se cortó por max_tokens antes de terminar. "
            "No se guardó en caché ni se consumió el cupo mensual - intenta de nuevo."
        )

    texto = "".join(block.text for block in response.content if block.type == "text")

    tokens_in = response.usage.input_tokens
    tokens_out = response.usage.output_tokens
    precio = PRECIO_POR_MILLON_USD[MODELO_ANALISIS]
    costo_usd = (tokens_in / 1_000_000 * precio["input"]) + (tokens_out / 1_000_000 * precio["output"])
    costo_cop = round(costo_usd * TRM_COP_POR_USD, 2)

    if cache:
        cache.huella_datos = huella
        cache.modelo = MODELO_ANALISIS
        cache.analisis_texto = texto
        cache.tokens_entrada = tokens_in
        cache.tokens_salida = tokens_out
        cache.costo_cop = costo_cop
    else:
        cache = AnalisisIACache(
            idcliente=idcliente,
            tipo_reporte=tipo_reporte,
            periodo_desde=periodo_desde,
            periodo_hasta=periodo_hasta,
            huella_datos=huella,
            modelo=MODELO_ANALISIS,
            analisis_texto=texto,
            tokens_entrada=tokens_in,
            tokens_salida=tokens_out,
            costo_cop=costo_cop,
        )
        db.session.add(cache)

    db.session.commit()
    _incrementar_uso_mensual(idcliente)

    return {
        "fuente": "nuevo",
        "analisis": texto,
        "modelo": MODELO_ANALISIS,
        "generado_en": datetime.utcnow().isoformat(),
        "uso_mensual": consultar_uso_mensual(idcliente),
        "tope_mensual": TOPE_MENSUAL,
    }


PROMPT_SISTEMA_DIAGNOSTICO_INTEGRAL = (
    "Eres un analista financiero senior / CFO fraccional entregando un "
    "diagnóstico COMPLETO e INTEGRADO de una PyME colombiana al dueño o "
    "gerente, combinando en un solo análisis cuatro fuentes: el Estado de "
    "Resultados (rentabilidad del período), el Balance General "
    "(estructura patrimonial al corte), los Indicadores Financieros "
    "(liquidez, endeudamiento, ROE/ROA, ciclo de caja) y el panel "
    "operativo del Resumen Ejecutivo (caja disponible, autonomía de caja, "
    "utilidad neta del período - la línea de fondo, lo que realmente le "
    "queda a la empresa -, egresos totales - costos y gastos reconocidos "
    "contablemente, la contraparte de las ventas -, evolución mensual de "
    "ventas/egresos/EBITDA/utilidad neta/eficiencia operativa, "
    "concentración en clientes/proveedores/gastos principales). Vas a "
    "recibir las cuatro ya calculadas en JSON - no inventes cifras que no "
    "estén ahí. Si alguna fuente viene marcada como no disponible "
    "(\"disponible\": false), dilo explícitamente y seguí con las demás, "
    "no la inventes.\n\n"
    "ESTO NO ES UN RESUMEN DE CADA REPORTE POR SEPARADO. El objetivo "
    "central es CONECTAR la información entre fuentes - encontrar lo que "
    "solo se ve combinándolas. Ejemplos del tipo de cruce esperado "
    "(usá los datos reales para encontrar los tuyos, no copies estos "
    "textualmente si no aplican):\n"
    "- Si la utilidad neta o el EBITDA del período son aceptables pero el "
    "Balance muestra que buena parte del patrimonio es 'calculado' (no "
    "contabilizado explícitamente), esa utilidad no se está reflejando "
    "correctamente en los libros - decilo.\n"
    "- Si hay una brecha grande entre EBITDA y utilidad neta (ej. EBITDA "
    "positivo pero utilidad neta negativa, o una caída de utilidad neta "
    "mucho mayor que la de EBITDA), buscá la explicación en gastos "
    "financieros o partidas no operacionales - cruzalo con la deuda "
    "financiera del Balance/Indicadores si aplica.\n"
    "- Si Indicadores muestra ROE alto, revisá si el Balance explica que "
    "es apalancamiento y no eficiencia real (ROA bajo o activo mínimo).\n"
    "- Si la autonomía de caja (runway) del panel operativo es corta, "
    "cruzala con el ciclo de cobro/pago (DSO/DPO de Indicadores) y con la "
    "cartera o cuentas por pagar más antiguas para explicar POR QUÉ es "
    "corta, no solo cuánto es.\n"
    "- Si la evolución mensual muestra un mes atípico (una caída o un "
    "pico que no se repite en los demás meses), señalalo explícitamente y "
    "conectalo con lo que el Balance o Indicadores puedan explicar de ese "
    "mismo período.\n\n"
    "Sobre la evolución mensual (evolucion_mensual del panel operativo): "
    "explicá la tendencia real de ventas, egresos totales, EBITDA, "
    "utilidad neta y eficiencia operativa a lo largo del período "
    "seleccionado, no solo el número del último corte - decí si mejora, "
    "empeora, es estable o errática mes a mes, y qué mes se sale del "
    "patrón y por qué importa. Ventas y egresos van de la mano: si un mes "
    "los egresos crecen más rápido que las ventas (o viceversa), decilo "
    "explícitamente.\n\n"
    "Responde siempre en español, en formato markdown, con esta "
    "estructura fija y sin exceder ~3000-3500 palabras en total (podés "
    "usar tablas markdown donde ayuden a mostrar un desglose):\n\n"
    "## Diagnóstico general\n"
    "2-4 frases: un veredicto claro de cómo está la empresa en conjunto "
    "hoy, sin rodeos.\n\n"
    "## Fortalezas\n"
    "Entre 2 y 4, cada una con el dato que la respalda.\n\n"
    "## Riesgos y focos de atención\n"
    "Entre 4 y 7, priorizados por impacto real para el negocio (no por "
    "el reporte del que vienen), cada uno con el número que lo respalda.\n"
    "\n"
    "## Conexiones entre reportes\n"
    "Entre 2 y 4 cruces explícitos como los del ejemplo de arriba - esta "
    "es la sección más importante, la que justifica un diagnóstico "
    "integral en vez de leer cada reporte por separado.\n\n"
    "## Recomendaciones priorizadas\n"
    "Entre 4 y 6 acciones concretas, en orden de urgencia/impacto, "
    "ligadas a los hallazgos de arriba, sin importar de qué fuente "
    "vengan.\n\n"
    "Si algo no se puede determinar con los datos disponibles, dilo "
    "explícitamente en vez de asumir u omitirlo."
)


def _resumen_dashboard_para_ia(dashboard_data: dict) -> dict:
    """Payload curado de la capa operativa exclusiva del Resumen
    Ejecutivo - caja, autonomia de caja, evolucion mensual, y
    concentracion en clientes/proveedores/gastos. Nada de esto vive en
    PyG/Balance/Indicadores, por eso se manda aparte.

    IMPORTANTE: NO se manda el dict "periodo" completo. Trae campos como
    ultima_fecha_auxiliar/fecha_corte_confiable que reflejan el estado
    del sistema "a hoy" (avanzan solos con cada sincronizacion nueva),
    sin relacion con el periodo Jan-Jul que se esta analizando - si
    entran a la huella, cualquier sincronizacion de rutina hace pensar
    al sistema que los datos del periodo cambiaron cuando en realidad
    siguen iguales, y dispara regeneraciones (con costo) innecesarias.
    Solo se incluyen los campos que describen el periodo en si."""
    periodo = dashboard_data.get("periodo", {})
    return {
        "periodo": {
            "desde": periodo.get("desde"),
            "hasta": periodo.get("hasta"),
            "anterior_desde": periodo.get("anterior_desde"),
            "anterior_hasta": periodo.get("anterior_hasta"),
            "modo_periodo": periodo.get("modo_periodo"),
            "tipo_corte": periodo.get("tipo_corte"),
        },
        "kpis_operativos": dashboard_data.get("kpis", {}),
        "evolucion_mensual": dashboard_data.get("series", {}).get("mensual", []),
        "top_gastos": dashboard_data.get("top_gastos", []),
        "top_clientes": dashboard_data.get("top_clientes", []),
        "top_proveedores": dashboard_data.get("top_proveedores", []),
        "explicaciones_sistema": dashboard_data.get("explicaciones", []),
        "acciones_sugeridas_sistema": dashboard_data.get("acciones", []),
        "alertas_sistema": dashboard_data.get("alertas", []),
    }


def _resumen_transversal_para_ia(
    pnl_data: dict, balance_data: dict, indicadores_data: dict, dashboard_data: dict
) -> dict:
    """Ensambla los 3 payloads ya validados (PyG/Balance/Indicadores) mas
    la capa operativa del dashboard - no reconstruye extraccion de datos,
    solo reusa lo que cada reporte individual ya probo. Si Balance o
    Indicadores no se pudieron construir para el periodo (ej. snapshot
    faltante), se marca como no disponible en vez de fallar todo el
    diagnostico - PyG y el panel operativo solos ya aportan valor."""
    return {
        "estado_resultados": _resumen_pyg_para_ia(pnl_data),
        "balance_general": (
            _resumen_balance_para_ia(balance_data)
            if balance_data.get("ok")
            else {"disponible": False, "motivo": balance_data.get("error", "No disponible para este período")}
        ),
        "indicadores_financieros": (
            _resumen_indicadores_para_ia(indicadores_data)
            if indicadores_data.get("ok")
            else {"disponible": False, "motivo": indicadores_data.get("error", "No disponible para este período")}
        ),
        "panel_operativo": _resumen_dashboard_para_ia(dashboard_data),
    }


def generar_analisis_transversal(
    idcliente: int,
    desde: str,
    hasta: str,
    pnl_data: dict,
    balance_data: dict,
    indicadores_data: dict,
    dashboard_data: dict,
    forzar: bool = False,
    solo_verificar: bool = False,
) -> dict:
    """
    Punto de entrada principal para el Diagnóstico Integral (Panel
    Ejecutivo). Mismo patron de cache/tope/huella que los otros 3.

    solo_verificar=True corta antes de llamar a la IA (ver
    generar_analisis_pyg).
    """
    tipo_reporte = "diagnostico_integral"
    periodo_desde = date.fromisoformat(desde)
    periodo_hasta = date.fromisoformat(hasta)

    payload = _resumen_transversal_para_ia(pnl_data, balance_data, indicadores_data, dashboard_data)
    huella = _huella_datos(payload)

    cache = AnalisisIACache.query.filter_by(
        idcliente=idcliente,
        tipo_reporte=tipo_reporte,
        periodo_desde=periodo_desde,
        periodo_hasta=periodo_hasta,
    ).first()

    if solo_verificar:
        return _resultado_verificacion(cache, huella)

    if cache and cache.huella_datos == huella and not forzar:
        return {
            "fuente": "cache",
            "analisis": cache.analisis_texto,
            "modelo": cache.modelo,
            "generado_en": cache.updated_at.isoformat() if cache.updated_at else None,
            "uso_mensual": consultar_uso_mensual(idcliente),
            "tope_mensual": TOPE_MENSUAL,
        }

    uso_actual = consultar_uso_mensual(idcliente)
    if uso_actual >= TOPE_MENSUAL:
        raise TopeAlcanzadoError(uso_actual, TOPE_MENSUAL)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY no está configurada en el entorno del backend.")

    client = anthropic.Anthropic(api_key=api_key)

    mensaje_usuario = (
        f"Aquí está el diagnóstico integral de la empresa para el período "
        f"{desde} a {hasta}, combinando Estado de Resultados, Balance "
        f"General, Indicadores Financieros y el panel operativo del "
        f"Resumen Ejecutivo, en formato JSON:\n\n"
        f"{json.dumps(payload, ensure_ascii=False, default=str)}"
    )

    with client.beta.messages.stream(
        model=MODELO_ANALISIS,
        max_tokens=32000,
        system=PROMPT_SISTEMA_DIAGNOSTICO_INTEGRAL,
        thinking={"type": "adaptive"},
        betas=["server-side-fallback-2026-07-01"],
        fallbacks="default",
        messages=[{"role": "user", "content": mensaje_usuario}],
    ) as stream:
        response = stream.get_final_message()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            "La respuesta de la IA se cortó por max_tokens antes de terminar. "
            "No se guardó en caché ni se consumió el cupo mensual - intenta de nuevo."
        )

    texto = "".join(block.text for block in response.content if block.type == "text")

    tokens_in = response.usage.input_tokens
    tokens_out = response.usage.output_tokens
    precio = PRECIO_POR_MILLON_USD[MODELO_ANALISIS]
    costo_usd = (tokens_in / 1_000_000 * precio["input"]) + (tokens_out / 1_000_000 * precio["output"])
    costo_cop = round(costo_usd * TRM_COP_POR_USD, 2)

    if cache:
        cache.huella_datos = huella
        cache.modelo = MODELO_ANALISIS
        cache.analisis_texto = texto
        cache.tokens_entrada = tokens_in
        cache.tokens_salida = tokens_out
        cache.costo_cop = costo_cop
    else:
        cache = AnalisisIACache(
            idcliente=idcliente,
            tipo_reporte=tipo_reporte,
            periodo_desde=periodo_desde,
            periodo_hasta=periodo_hasta,
            huella_datos=huella,
            modelo=MODELO_ANALISIS,
            analisis_texto=texto,
            tokens_entrada=tokens_in,
            tokens_salida=tokens_out,
            costo_cop=costo_cop,
        )
        db.session.add(cache)

    db.session.commit()
    _incrementar_uso_mensual(idcliente)

    return {
        "fuente": "nuevo",
        "analisis": texto,
        "modelo": MODELO_ANALISIS,
        "generado_en": datetime.utcnow().isoformat(),
        "uso_mensual": consultar_uso_mensual(idcliente),
        "tope_mensual": TOPE_MENSUAL,
    }


# ─────────────────────────────────────────────
# Exportar a Word
# ─────────────────────────────────────────────
# Generado en el backend con python-docx, NO en el navegador. La libreria
# JS `docx` se probo primero del lado del cliente y resulto incompatible
# con el bundler de Next.js (Turbopack) - su propio codigo empaquetado
# usa un patron de transpilacion de clases (super() dentro de una
# funcion anidada) que Turbopack no logra procesar, rompiendo el chunk
# entero. Generarlo aca evita el problema de raiz y es mas simple.

LOGO_PATH = os.path.join(os.path.dirname(__file__), "static", "branding", "insightsflow-logo.png")


def _agregar_texto_con_negritas(paragraph, texto: str) -> None:
    partes = re.split(r"(\*\*.+?\*\*)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        else:
            paragraph.add_run(parte)


def _es_separador_tabla(linea: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$", linea.strip()))


def _dividir_fila(linea: str) -> list[str]:
    t = linea.strip()
    if t.startswith("|"):
        t = t[1:]
    if t.endswith("|"):
        t = t[:-1]
    return [c.strip() for c in t.split("|")]


def _marcar_fila_como_encabezado_repetible(row) -> None:
    """Hace que la fila de encabezado de una tabla se repita en cada hoja
    cuando la tabla se parte entre páginas - python-docx no expone esto
    como propiedad de alto nivel, hay que agregar w:tblHeader al XML de
    la fila directamente. Sin esto (visto en un PDF real de Diagnóstico
    Integral), una tabla larga que cruza de página 1 a 2 arranca la
    página 2 con puros números, sin decir qué es cada columna."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _markdown_a_docx(doc: DocxDocument, markdown: str) -> None:
    # Los blockquotes (líneas "> texto", que Opus usa para notas
    # metodológicas tipo "> **Nota:** ...") no tenían manejo propio y
    # caían al párrafo genérico con el ">" literal sin quitar - se
    # despoja aca antes de parsear, el resto de la línea sigue su
    # camino normal (incluido el negritas).
    lineas = [
        re.sub(r"^\s*>\s?", "", l) if l.strip().startswith(">") else l
        for l in markdown.split("\n")
    ]
    i = 0
    n = len(lineas)

    while i < n:
        linea = lineas[i]

        if not linea.strip():
            i += 1
            continue

        encabezado = re.match(r"^(#{1,4})\s+(.*)$", linea)
        if encabezado:
            nivel = min(len(encabezado.group(1)), 3)
            doc.add_heading(encabezado.group(2), level=nivel)
            i += 1
            continue

        if "|" in linea and i + 1 < n and _es_separador_tabla(lineas[i + 1]):
            celdas_header = _dividir_fila(linea)
            j = i + 2
            filas = []
            while j < n and "|" in lineas[j] and lineas[j].strip():
                filas.append(_dividir_fila(lineas[j]))
                j += 1

            tabla = doc.add_table(rows=1, cols=len(celdas_header))
            try:
                tabla.style = "Light Grid Accent 1"
            except KeyError:
                pass
            hdr_cells = tabla.rows[0].cells
            for idx, texto_celda in enumerate(celdas_header):
                run = hdr_cells[idx].paragraphs[0].add_run(texto_celda)
                run.bold = True
            _marcar_fila_como_encabezado_repetible(tabla.rows[0])
            for fila in filas:
                cells = tabla.add_row().cells
                for idx, texto_celda in enumerate(fila):
                    if idx < len(cells):
                        # No usar cells[idx].text = texto_celda: eso vuelca
                        # el string crudo sin procesar markdown, dejando
                        # "**subtotal**" literal en vez de negrita - mismo
                        # helper que ya usan párrafos y listas.
                        _agregar_texto_con_negritas(cells[idx].paragraphs[0], texto_celda)
            doc.add_paragraph("")
            i = j
            continue

        if re.match(r"^-{3,}$", linea.strip()):
            i += 1
            continue

        if re.match(r"^-\s+", linea.strip()):
            while i < n and re.match(r"^-\s+", lineas[i].strip()):
                p = doc.add_paragraph(style="List Bullet")
                _agregar_texto_con_negritas(p, re.sub(r"^-\s+", "", lineas[i].strip()))
                i += 1
            continue

        if re.match(r"^\d+\.\s+", linea.strip()):
            while i < n and re.match(r"^\d+\.\s+", lineas[i].strip()):
                p = doc.add_paragraph(style="List Number")
                _agregar_texto_con_negritas(p, re.sub(r"^\d+\.\s+", "", lineas[i].strip()))
                i += 1
            continue

        lineas_parrafo = []
        while (
            i < n
            and lineas[i].strip() != ""
            and not re.match(r"^#{1,4}\s", lineas[i])
            and not re.match(r"^-{3,}$", lineas[i].strip())
            and not re.match(r"^-\s+", lineas[i].strip())
            and not re.match(r"^\d+\.\s+", lineas[i].strip())
            and not ("|" in lineas[i] and i + 1 < n and _es_separador_tabla(lineas[i + 1]))
        ):
            lineas_parrafo.append(lineas[i])
            i += 1
        p = doc.add_paragraph()
        _agregar_texto_con_negritas(p, " ".join(lineas_parrafo))


def _abreviar_valor_chart(valor: float) -> str:
    """Version compacta (sin decimales) de abreviar() del frontend -
    pensada solo para las etiquetas del grafico estatico, donde el
    espacio es angosto y el texto largo (con decimales) se solapa entre
    barras vecinas."""
    n = float(valor or 0)
    abs_n = abs(n)
    if abs_n >= 1_000_000:
        return f"{round(n / 1_000_000):,}M".replace(",", ".")
    if abs_n >= 1_000:
        return f"{round(n / 1_000):,}K".replace(",", ".")
    return f"{round(n):,}".replace(",", ".")


def generar_grafico_evolucion(evolucion: list) -> BytesIO | None:
    """Genera el mismo grafico de tendencia (Ingresos/Costos y Gastos/
    EBITDA) que se ve en pantalla, como PNG estatico, para insertarlo en
    el Word - reusa la data de construir_pnl(), no le cuesta nada a la
    IA. En pantalla el grafico no tiene eje Y, pero SI muestra el valor
    encima de cada barra/punto (LabelList) - sin eso, un grafico sin eje
    no se puede leer. Replicado aca con las mismas etiquetas."""
    if not evolucion:
        return None

    labels = [e.get("label", "") for e in evolucion]
    ingresos = [float(e.get("ingresos_totales") or 0) for e in evolucion]
    costos_gastos = [float(e.get("costos_gastos") or 0) for e in evolucion]
    ebitda = [float(e.get("ebitda") or 0) for e in evolucion]

    x = list(range(len(labels)))
    ancho = 0.32
    valor_max = max(ingresos + costos_gastos + [abs(v) for v in ebitda] + [1])
    valor_min_ebitda = min(ebitda + [0])

    # Mas ancho del que se ve en pantalla a proposito: en un documento
    # estatico (sin tooltip/hover) el unico recurso para que 7 grupos de
    # barras con etiquetas no se encimen es darles mas aire horizontal.
    fig, ax = plt.subplots(figsize=(9.6, 3.2), dpi=160)
    caja_etiqueta = dict(boxstyle="round,pad=0.12", fc="white", ec="none", alpha=0.88)

    barras_ingresos = ax.bar([i - ancho / 2 for i in x], ingresos, width=ancho, color=COLOR_INGRESOS, label="Ingresos")
    barras_costos = ax.bar([i + ancho / 2 for i in x], costos_gastos, width=ancho, color=COLOR_COSTOS_GASTOS, label="Costos y Gastos")
    ax.plot(x, ebitda, color=COLOR_EBITDA, linewidth=2.5, marker="o", markersize=4, label="EBITDA", zorder=5)

    offset = valor_max * 0.05
    for barra in list(barras_ingresos) + list(barras_costos):
        alto = barra.get_height()
        ax.text(
            barra.get_x() + barra.get_width() / 2, alto + offset, _abreviar_valor_chart(alto),
            ha="center", va="bottom", fontsize=7, fontweight="bold", color="#334155", bbox=caja_etiqueta,
        )
    for xi, yi in zip(x, ebitda):
        y_texto = yi + offset if yi >= 0 else yi - offset
        ax.text(
            xi, y_texto, _abreviar_valor_chart(yi),
            ha="center", va="bottom" if yi >= 0 else "top", fontsize=7, fontweight="bold",
            color=COLOR_EBITDA, bbox=caja_etiqueta, zorder=6,
        )

    # Margen inferior generoso y desacoplado del offset de las etiquetas
    # (con el offset solo, un EBITDA muy negativo terminaba con su
    # etiqueta pisando el nombre del mes en el eje X).
    limite_inferior = min(valor_min_ebitda * 2.3, -valor_max * 0.08)
    ax.set_ylim(top=valor_max * 1.32, bottom=limite_inferior)
    ax.set_xlim(-0.6, len(labels) - 0.4)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8, fontweight="bold")
    ax.get_yaxis().set_visible(False)
    for spine in ("top", "right", "left"):
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#cbd5e1")
    ax.tick_params(axis="x", length=0)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.22), ncol=3, frameon=False, fontsize=8)

    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def generar_word_analisis(
    markdown: str, nombre_cliente: str, periodo: str, evolucion: list | None = None
) -> BytesIO:
    doc = DocxDocument()

    seccion = doc.sections[0]

    header = seccion.header
    header_parrafo = header.paragraphs[0]
    if os.path.exists(LOGO_PATH):
        header_parrafo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    else:
        run_logo = header_parrafo.add_run("InsightsFlow")
        run_logo.bold = True

    info_parrafo = header.add_paragraph()
    info_run = info_parrafo.add_run(f"{nombre_cliente}  ·  {periodo}")
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    footer = seccion.footer
    footer_parrafo = footer.paragraphs[0]
    footer_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_parrafo.add_run(
        f"Reporte generado por la IA de InsightsFlow {datetime.utcnow().year}"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    grafico = generar_grafico_evolucion(evolucion) if evolucion else None
    if grafico:
        doc.add_picture(grafico, width=Inches(6.2))
        doc.add_paragraph("")

    _markdown_a_docx(doc, markdown)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def listar_historial_analisis(idcliente: int, tipo_reporte: str = "estado_resultados", limite: int = 24) -> list:
    """Periodos que este cliente ya tiene analizados (en analisis_ia_cache)
    para que el frontend pueda ofrecer 'volver a ver' sin gastar cupo -
    revisitar un periodo ya analizado siempre sale del cache, gratis."""
    filas = (
        AnalisisIACache.query
        .filter_by(idcliente=idcliente, tipo_reporte=tipo_reporte)
        .order_by(AnalisisIACache.periodo_desde.desc())
        .limit(limite)
        .all()
    )
    return [
        {
            "periodo_desde": f.periodo_desde.isoformat(),
            "periodo_hasta": f.periodo_hasta.isoformat(),
            "generado_en": f.updated_at.isoformat() if f.updated_at else None,
            "modelo": f.modelo,
        }
        for f in filas
    ]
